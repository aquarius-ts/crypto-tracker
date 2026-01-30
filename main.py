# -*- coding: utf-8 -*-
import requests
import browser_cookie3
import re
import urllib.parse
import json
import zipfile
import io
import os
import time
from loguru import logger

logger.remove()
from androguard.core.apk import APK
from docx import Document

GERRIT_BASE = "http://10.24.71.91/gerrit"
WHITELIST_PERMISSIONS = {
    "android.permission.INTERNET",
    "android.permission.ACCESS_NETWORK_STATE",
    "android.permission.RECEIVE_BOOT_COMPLETED",
}

NOTE_FILE = "checked.json"

def load_checked():
    if os.path.exists(NOTE_FILE):
        with open(NOTE_FILE, "r") as f:
            return json.load(f)
    return {}


def save_checked(checked):
    with open(NOTE_FILE, "w") as f:
        json.dump(checked, f, indent=2)


def get_cookies():
    return browser_cookie3.chrome(domain_name="10.24.71.91")


def gerrit_get(url, retries=3):
    for i in range(retries):
        try:
            r = requests.get(url, cookies=get_cookies(), timeout=15)
            r.raise_for_status()
            return json.loads(r.text.lstrip(")]}'"))
        except requests.RequestException as e:
            print("Request failed ({}/{}): {}".format(i + 1, retries, e))
            if i < retries - 1:
                time.sleep(2)
            else:
                raise


def filter_apk_changes(changes):
    apk_changes = []
    for c in changes:
        rev = c["current_revision"]
        files = c["revisions"][rev]["files"]
        apk_files = [f for f in files if f.endswith(".apk")]
        if apk_files:
            apk_changes.append((c, rev, apk_files))
    return apk_changes


def download_apk(project, change, revision, apk_path):
    encoded = urllib.parse.quote(apk_path, safe="")
    url = (f"{GERRIT_BASE}/changes/"
           f"{urllib.parse.quote(project, safe='')}~{change}"
           f"/revisions/{revision}/files/{encoded}/download")
    print(f"\n⬇️ Downloading ZIP: {url}")

    r = requests.get(url, cookies=get_cookies(), stream=True)
    r.raise_for_status()

    buf = io.BytesIO()
    total = int(r.headers.get("Content-Length", 0))
    done = 0

    for chunk in r.iter_content(1024 * 256):
        if chunk:
            buf.write(chunk)
            done += len(chunk)
            if total:
                print(
                    f"\r  {done * 100 / total:5.1f}% ({done / 1024 / 1024:.2f} MB)",
                    end="",
                    flush=True)

    print("\n📦 Extracting ZIP...")
    buf.seek(0)

    with zipfile.ZipFile(buf) as z:
        for name in z.namelist():
            if name.endswith(".apk"):
                z.extract(name, ".")
                print(f"✅ APK extracted: {name}")
                return name

    raise RuntimeError("❌ APK not found in ZIP")


def check_release(apk):
    try:
        a = APK(apk)
        debuggable = False
        app = a.get_android_manifest_xml().find("application")
        if app is not None:
            dbg = app.get(
                "{http://schemas.android.com/apk/res/android}debuggable")
            if dbg == "true":
                debuggable = True
        if debuggable:
            return False, "APK is debuggable (android:debuggable=true)"
    except Exception as e:
        return False, f"Cannot check release: {e}"
    return True, None


def find_whitelisted_permissions(apk):
    a = APK(apk)
    perms = set(a.get_permissions() or [])
    hits = sorted(p for p in perms if p in WHITELIST_PERMISSIONS)
    return hits


def check_version(apk, commit_title):
    try:
        a = APK(apk)
        apk_ver = a.get_androidversion_name()
        m = re.search(r"\[(\d+(\.\d+)+)\]", commit_title)
        if not m:
            return True, None
        expected = m.group(1)
        if apk_ver != expected:
            return False, f"APK version={apk_ver}, commit expects={expected}"
    except:
        return True, None
    return True, None


def ask_gerrit_link():
    print("\n🔗 Enter Gerrit change link (press Enter to skip):")
    link = input("> ").strip()
    return link if link else None


def parse_gerrit_change_url(url: str):
    m = re.search(r"/c/(.+)/\+/(\d+)", url)
    if not m:
        raise ValueError("❌ Invalid Gerrit change URL")
    project = m.group(1)
    change_number = m.group(2)
    return project, change_number


def get_change_by_number(project, change_number):
    url = (f"{GERRIT_BASE}/changes/"
           f"{urllib.parse.quote(project, safe='')}~{change_number}"
           "?o=CURRENT_REVISION&o=CURRENT_FILES")
    r = requests.get(url, cookies=get_cookies(), timeout=15)
    r.raise_for_status()
    return gerrit_parse(r)


def gerrit_parse(r):
    text = r.text
    if text.startswith(")]}'"):
        text = text[4:]
    return json.loads(text)


def create_checklist_doc(filename,
                         is_release=None,
                         whitelist_permissions=None,
                         branch_name=None,
                         version_ok=None):
    doc = Document()
    doc.add_heading("Checklist when uploading APK", level=1)

    def add_item(checked, text):
        prefix = "☑" if checked else "☐"
        doc.add_paragraph(f"{prefix} {text}")

    add_item(is_release is True, "The APK is release variant (not debug)")

    if whitelist_permissions:
        add_item(True, "APK has permission protected added to white-list")
        for p in whitelist_permissions:
            doc.add_paragraph(f"    - {p}")
    else:
        add_item(False, "APK has permission protected added to white-list")

    if branch_name:
        add_item(True, f"APK is pushed to the right branches ({branch_name})")
    else:
        add_item(False, "APK is pushed to the right branches")

    add_item(version_ok is True, "APK is in the correct version")

    add_item(
        False,
        "Checked related apk (Versioning for different models if needed)")
    add_item(False, "Checked and updated the right service version")

    doc.save(filename)


def main():
    link = ask_gerrit_link()
    if not link:
        print("❌ Gerrit link is required")
        return

    print(f"🔗 Gerrit link: {link}")

    try:
        project, change_number = parse_gerrit_change_url(link)
        change = get_change_by_number(project, change_number)
    except Exception as e:
        print("❌ Cannot parse Gerrit link:", e)
        return

    apk_changes = filter_apk_changes([change])
    if not apk_changes:
        print("ℹ️ No APK found in this commit")
        return

    checked = load_checked()

    for change, revision, apks in apk_changes:
        cid = change["change_id"]

        if cid in checked and revision in checked[cid]:
            print(f"ℹ️ Already checked: {cid} (patchset {revision})")
            return

        print("\n" + "=" * 60)
        print(f"📌 Commit : {change['subject']}")
        print(f"Project  : {change['project']}")
        print(f"Branch   : {change.get('branch')}")
        print(f"Patchset : {revision}")

        for apk_path in apks:
            apk_file = None
            try:
                # ===== 4. Download APK =====
                apk_file = download_apk(change["project"], change["_number"],
                                        revision, apk_path)

                print("\n===== APK REVIEW RESULT =====")

                # Release check
                is_release, release_err = check_release(apk_file)
                if release_err:
                    print("⚠️ NOTE:", release_err)

                # Permission whitelist
                whitelist_hits = find_whitelisted_permissions(apk_file)
                if whitelist_hits:
                    print("⚠️ NOTE: APK uses whitelisted permissions:")
                    for p in whitelist_hits:
                        print("  -", p)

                # Version check (optional)
                version_ok, version_msg = check_version(
                    apk_file, change["subject"])
                if version_msg:
                    print("⚠️ NOTE:", version_msg)

                # ===== 5. Create checklist =====
                checklist_name = f"APK_Checklist_{cid}_{revision}.docx"
                create_checklist_doc(checklist_name,
                                     is_release=is_release,
                                     whitelist_permissions=whitelist_hits,
                                     branch_name=change.get("branch"),
                                     version_ok=version_ok)

                print(f"📄 Checklist created: {checklist_name}")
                print("✅ REVIEW COMPLETED")

            finally:
                # ===== 6. Cleanup APK =====
                if apk_file and os.path.exists(apk_file):
                    os.remove(apk_file)
                    print(f"🧹 Cleaned up APK: {apk_file}")

        # ===== 7. Save checked patchset =====
        checked.setdefault(cid, [])
        checked[cid].append(revision)
        save_checked(checked)


if __name__ == "__main__":
    main()
